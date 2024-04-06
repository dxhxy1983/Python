import os,sys
import pandas as pd
import shutil
import re
from docx import Document
from docx.shared import Pt
def open_excel(path):
    try:    
        book = pd.read_excel(path)
        return book
    except:
        print("open excel file failed!")

def find_files(folder, extension):
    file_list = []
    for root, dirs, files in os.walk(folder):
        for file in files:
            if file.endswith(extension):
                file_list.append(os.path.join(root, file))
    return file_list
def add_formatted_paragraph(doc, text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(12)  # 设置段落字号为12
    return paragraph
def save_as_doc(current_path,file_path):
    df=pd.read_excel(file_path)
    a=[]
    a=file_path.split("\\")
    file_name,file_style=a[-1].split(".")
    # 创建一个新的Word文档
    word_file_path =current_path+"\\"+file_name+ '.docx'
    doc = Document()
    #增加标题
    doc.add_heading(file_name, level=1)
    s={}
    s[0]="一、"+str(df.iloc[1,1])+str(df.iloc[1,2])+"元,其中:"
    for i in range(2,6):
        for j in range(1,4):
            s[0]=s[0]+str(df.iloc[i,2*j-1])+str(df.iloc[i,2*j])+"元,"
    s[0]=s[0]+str(df.iloc[7,1])+str(df.iloc[7,2])+"元。"
    
    s[1]="二、"+str(df.iloc[8,1])+str(df.iloc[8,2])+"元,"+str(df.iloc[8,3])+str(df.iloc[8,4])+"元。"
    s[2]="三、"+str(df.iloc[10,1])+str(df.iloc[10,2])+"元,"+str(df.iloc[10,3])+str(df.iloc[10,4])+"元,故本月"+str(df.iloc[9,1])+str(df.iloc[9,2])+"元。"
    s[3]="四、"+str(df.iloc[11,1])+str(df.iloc[11,2])+"元,"+str(df.iloc[11,4])+str("{:.2%}".format(df.iloc[11,5]))
    s[4]="五、本月实际产生"+str(df.iloc[13,1])+str(df.iloc[13,2])+"元,"+str(df.iloc[13,3])+str(df.iloc[13,4])+"元,"+str(df.iloc[13,5])+str(df.iloc[13,6])+"元,合计"+str(df.iloc[12,1])+str(df.iloc[12,2])+"元。"
    s[5]="六、"+str(df.iloc[14,1])+str(df.iloc[14,2])+"元。"
    s[6]="七、"+str(df.iloc[15,1])+str(df.iloc[15,2])+"元。"
    for i in range(len(s)):
        add_formatted_paragraph(doc,s[i])
    doc.save(word_file_path)


if __name__=="__main__":
    #获取当前文件目录
    current_path = os.path.dirname(os.path.realpath(sys.argv[0]))
    
    extension=".xlsx"
    file_list = find_files(current_path, extension)
    # print(file_list)
    # print(current_path)
    if file_list!=None:
        for i in range(len(file_list)):
            file_path=file_list[i]
            save_as_doc(current_path,file_path)

    # print(f'文件已保存至：{word_file_path}')

    
    